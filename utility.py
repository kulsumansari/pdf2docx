import requests
from pdf2docx import Converter
from docx import Document
import pdfplumber


def download_pdf(pdf_url, local_pdf_path):
    
    response = requests.get(pdf_url)
    with open(local_pdf_path, 'wb') as file:
        file.write(response.content)
    print(f"Downloaded PDF from {pdf_url} to {local_pdf_path}")


def extract_text_from_pdf(pdf_url, local_pdf_path):
    response = requests.get(pdf_url)
    with open(local_pdf_path, 'wb') as file:
        file.write(response.content)
    
    text = ""
    paragraphs = []
    with pdfplumber.open(local_pdf_path) as pdf:
        header_footer_height=40

        for page in pdf.pages:
            # width = page.width
            height = page.height
            text = ""
            
            # Extract text with words and handle links
            for element in page.extract_words():
                # Skip elements in the header or footer area
                if element['top'] > header_footer_height and element['bottom'] < height - header_footer_height:
                    text += f"{element['text']} "
            

            if text:
                # Split the text by double newline to get paragraphs
                page_paragraphs = text.split('\n\n')
                paragraphs.extend(page_paragraphs)

    # for i, paragraph in enumerate(paragraphs):
    #     print(f"Paragraph {i + 1}:\n{paragraph}\n")

    return paragraphs
    

# convert to local docx
def convertPdf2Docx(local_pdf_path, local_docx_path):

    cv = Converter(local_pdf_path)
    cv.convert(local_docx_path)
    cv.close()
    print(f"Converted PDF from {local_pdf_path} to DOCX at {local_docx_path}")
    


def convert_json_to_docx(content,  document=None):

    if document is None:
        document = Document()
    
    # if condition to handle if content is an array of object , and object can have any key
    if isinstance(content, list):
        for item in content:
            if isinstance(item, dict):
                for key, value in item.items():
                    if key == 'title':
                        document.add_heading(str(value), 0)
                    elif key == 'heading':
                        document.add_heading(str(value), 1)
                    else:
                        document.add_paragraph(str(value))
    
    elif isinstance(content, dict):
        for key, value in content.items():
            if (key == 'title'):
                document.add_heading(str(value), 0)
            elif key == 'heading':
                document.add_heading(str(value), 1)

            elif isinstance(value, str):
                document.add_paragraph(str(value))

            elif isinstance(value, list):
                print(value)
                for item in value:
                    if isinstance(item, dict):
                        convert_json_to_docx(item, document)
    else:
        document.add_paragraph(str(item))
                

    return document


def save_document(document, output_docx):
    document.save(output_docx)
    print(f"Converted JSON to DOCX at {output_docx}")
        

# upload to CMS
def upload2CMS(local_docx_path, asset):
    import os

    title = asset['title']
    filename = asset['filename'].split('.')[0] + '.docx'
    print ('========' + filename + '=======')

    url = "https://api.contentstack.io/v3/assets"

    payload = {'asset[title]': title}
    files=[
        ('asset[upload]',(filename ,open(local_docx_path,'rb'),'application/octet-stream'))
    ]
    headers = {
        'api_key': os.environ.get('API_KEY'),
        'authorization': os.environ.get('AUTHORIZATION')
    }

    response = requests.request('POST', url, headers=headers, data=payload, files=files)

    return response